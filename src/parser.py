"""
parser.py
---------
Handles reading resumes in different file formats (PDF, DOCX, TXT) and
returning plain text so the rest of the pipeline never has to care about
file format again.
"""

import os


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    import pdfplumber

    text_chunks = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    import docx

    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of tables (some resumes use table layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(file_path: str) -> str:
    """
    Dispatch to the correct extractor based on file extension.
    Raises ValueError for unsupported formats.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in (".txt", ".md"):
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(
            f"Unsupported resume format '{ext}' for file: {file_path}. "
            f"Supported formats: .pdf, .docx, .txt"
        )


def load_resumes(folder_path: str) -> dict:
    """
    Load every supported resume file in a folder.
    Returns: { filename: raw_text }
    """
    supported_ext = (".pdf", ".docx", ".txt", ".md")
    resumes = {}

    for filename in sorted(os.listdir(folder_path)):
        if filename.lower().endswith(supported_ext):
            full_path = os.path.join(folder_path, filename)
            try:
                text = extract_text(full_path)
                if text.strip():
                    resumes[filename] = text
                else:
                    print(f"  [warn] {filename} produced no extractable text, skipping.")
            except Exception as e:
                print(f"  [warn] Failed to parse {filename}: {e}")

    return resumes
