"""
tests/test_parser.py
---------------------
Unit tests for src/parser.py -- reading resumes from PDF/DOCX/TXT files.
"""

import os
import pytest

from src.parser import (
    extract_text_from_txt,
    extract_text_from_pdf,
    extract_text_from_docx,
    extract_text,
    load_resumes,
)


@pytest.fixture
def sample_txt_file(tmp_path):
    content = "John Doe\nPython Developer\n5 years of experience"
    file_path = tmp_path / "resume.txt"
    file_path.write_text(content, encoding="utf-8")
    return str(file_path), content


@pytest.fixture
def sample_pdf_file(tmp_path):
    """Generate a tiny real PDF using reportlab so we test actual PDF parsing."""
    from reportlab.pdfgen import canvas

    file_path = tmp_path / "resume.pdf"
    c = canvas.Canvas(str(file_path))
    c.drawString(100, 750, "Jane Smith")
    c.drawString(100, 730, "Backend Engineer with 3 years of experience")
    c.save()
    return str(file_path)


@pytest.fixture
def sample_docx_file(tmp_path):
    import docx

    file_path = tmp_path / "resume.docx"
    doc = docx.Document()
    doc.add_paragraph("Alex Rivera")
    doc.add_paragraph("Full Stack Developer")
    doc.save(str(file_path))
    return str(file_path)


def test_extract_text_from_txt(sample_txt_file):
    path, expected_content = sample_txt_file
    result = extract_text_from_txt(path)
    assert result == expected_content


def test_extract_text_from_pdf(sample_pdf_file):
    result = extract_text_from_pdf(sample_pdf_file)
    assert "Jane Smith" in result
    assert "Backend Engineer" in result


def test_extract_text_from_docx(sample_docx_file):
    result = extract_text_from_docx(sample_docx_file)
    assert "Alex Rivera" in result
    assert "Full Stack Developer" in result


def test_extract_text_dispatches_by_extension(sample_txt_file):
    path, expected_content = sample_txt_file
    result = extract_text(path)
    assert result == expected_content


def test_extract_text_unsupported_format_raises(tmp_path):
    bad_file = tmp_path / "resume.xlsx"
    bad_file.write_text("some content")
    with pytest.raises(ValueError):
        extract_text(str(bad_file))


def test_load_resumes_reads_all_supported_files(tmp_path):
    (tmp_path / "a.txt").write_text("Resume A content", encoding="utf-8")
    (tmp_path / "b.txt").write_text("Resume B content", encoding="utf-8")
    (tmp_path / "ignore_me.xlsx").write_text("not supported")

    resumes = load_resumes(str(tmp_path))

    assert "a.txt" in resumes
    assert "b.txt" in resumes
    assert "ignore_me.xlsx" not in resumes
    assert resumes["a.txt"] == "Resume A content"


def test_load_resumes_skips_empty_files(tmp_path):
    (tmp_path / "empty.txt").write_text("", encoding="utf-8")
    (tmp_path / "real.txt").write_text("Has content", encoding="utf-8")

    resumes = load_resumes(str(tmp_path))

    assert "empty.txt" not in resumes
    assert "real.txt" in resumes


def test_load_resumes_on_empty_folder_returns_empty_dict(tmp_path):
    resumes = load_resumes(str(tmp_path))
    assert resumes == {}
